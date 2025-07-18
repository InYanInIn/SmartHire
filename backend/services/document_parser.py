import pypdfium2 as pdfium
from docx import Document
def extract_text(path: str) -> str:

    if path.endswith(".pdf"):
        pdf = pdfium.PdfDocument(path)
        text_content = []

        for page in pdf:
            textpage = page.get_textpage()
            text_content.append(textpage.get_text_range())

        return "\n".join(text_content)
    elif path.endswith(".docx"):
        doc = Document(path)
        return "\n".join([para.text for para in doc.paragraphs])
    elif path.endswith(".txt"):
        with open(path, 'r', encoding='utf-8') as file:
            return file.read()
    else:
        raise ValueError("Unsupported file format. Supported formats are: PDF, DOCX, TXT.")


